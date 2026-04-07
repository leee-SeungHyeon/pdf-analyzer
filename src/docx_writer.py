from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Optional
import os
import re


def write(analysis: Dict, extracted: Dict, output_path: str, source_filename: str = ""):
    """
    분석 결과와 추출 데이터를 docx 파일로 저장한다.

    섹션별로 LLM 분석 텍스트 → 원본 테이블 → 원본 이미지 순으로 렌더링한다.
    """
    doc = Document()
    _set_default_font(doc)

    # 제목
    title_text = "문서 분석 리포트"
    if source_filename:
        title_text += f": {os.path.basename(source_filename)}"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 문서 유형
    p = doc.add_paragraph()
    run = p.add_run(f"문서 유형: {analysis.get('doc_type', '알 수 없음')}")
    run.bold = True
    run.font.size = Pt(13)
    doc.add_paragraph()

    # 전체 요약
    doc.add_heading("전체 요약", level=1)
    doc.add_paragraph(analysis.get("overall_summary", "요약 없음"))
    doc.add_paragraph()

    # 핵심 인사이트
    key_insights = analysis.get("key_insights", [])
    if key_insights:
        doc.add_heading("핵심 인사이트", level=1)
        for insight in key_insights:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(insight)
        doc.add_paragraph()

    # 섹션별 분석 (LLM 줄글 + 원본 테이블 + 원본 이미지)
    llm_sections = analysis.get("sections", [])
    doc_sections = extracted.get("doc_sections", [])

    if llm_sections:
        doc.add_heading("섹션별 분석", level=1)
        for llm_sec in llm_sections:
            sec_title = llm_sec.get("title", "")
            sec_summary = llm_sec.get("summary", "")

            doc.add_heading(sec_title, level=2)

            if sec_summary:
                doc.add_paragraph(sec_summary)

            # 원문 섹션에서 테이블/이미지 찾아서 삽입
            doc_sec = _find_doc_section(sec_title, doc_sections)
            if doc_sec:
                for table_str in doc_sec.get("tables", []):
                    doc.add_paragraph()
                    _render_table(doc, table_str)

                for img_path in doc_sec.get("images", []):
                    doc.add_paragraph()
                    _insert_image(doc, img_path)

            doc.add_paragraph()

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    print(f"저장 완료: {output_path}")


def _find_doc_section(llm_title: str, doc_sections: List[Dict]) -> Optional[Dict]:
    """LLM 섹션 제목과 가장 잘 맞는 원문 섹션을 반환한다."""
    # 정확히 일치
    for ds in doc_sections:
        if ds["title"] == llm_title:
            return ds
    # 부분 일치
    llm_lower = llm_title.lower()
    for ds in doc_sections:
        ds_lower = ds["title"].lower()
        if llm_lower in ds_lower or ds_lower in llm_lower:
            return ds
    return None


def _render_table(doc: Document, table_str: str):
    """마크다운 테이블을 Word 테이블로 렌더링한다."""
    lines = [l for l in table_str.split("\n") if l.strip()]
    # 구분선 제거 (|---|---|)
    data_lines = [l for l in lines if not re.match(r'^\|[-:\s|]+\|$', l)]

    if not data_lines:
        return

    parsed_rows = []
    for row in data_lines:
        cells = [c.strip() for c in row.strip("|").split("|")]
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    max_cols = max(len(r) for r in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(parsed_rows):
        for c_idx in range(max_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = cell_text
            if r_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def _insert_image(doc: Document, image_path: str, max_width: float = 5.5):
    """이미지를 문서에 삽입한다."""
    try:
        doc.add_picture(image_path, width=Inches(max_width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"  이미지 삽입 실패 ({os.path.basename(image_path)}): {e}")


def _set_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
